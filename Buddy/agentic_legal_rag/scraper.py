#!/usr/bin/env python3
"""
Legal Data Scraper Module
Handles downloading and fetching data from various legal sources
"""

import os
import sys
import json
import requests
import zipfile
import pandas as pd
from typing import List, Dict, Any, Optional, Generator
from pathlib import Path
import logging
from dataclasses import dataclass
import time
import re
from bs4 import BeautifulSoup
import PyPDF2
import docx
from io import BytesIO
import subprocess
import kaggle
from datasets import load_dataset
import urllib.parse
from datetime import datetime

# Add current directory to path for imports
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from config import config

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class ScrapedDocument:
    """Raw scraped document structure"""
    title: str
    content: str
    source_url: str
    source_type: str  # 'indiacode', 'constitution', 'github', 'zenodo', 'ildc', 'kaggle', 'huggingface', 'lawsum'
    section_article: Optional[str] = None
    date: Optional[str] = None
    doc_type: Optional[str] = None
    metadata: Dict[str, Any] = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}

class IndiaCodeScraper:
    """Scraper for IndiaCode (https://www.indiacode.nic.in/)"""
    
    def __init__(self):
        self.base_url = "https://www.indiacode.nic.in"
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })
    
    def scrape_acts(self, max_acts: int = 50) -> List[ScrapedDocument]:
        """Scrape central acts from IndiaCode"""
        documents = []
        
        try:
            # Get list of acts
            acts_url = f"{self.base_url}/show-data?actid=AC_CEN_0001_00000001"
            response = self.session.get(acts_url)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Find act links (this is a simplified approach - actual implementation would need more sophisticated parsing)
            act_links = soup.find_all('a', href=True)
            
            count = 0
            for link in act_links[:max_acts]:
                if count >= max_acts:
                    break
                    
                href = link.get('href')
                if href and 'actid' in href:
                    try:
                        act_url = urllib.parse.urljoin(self.base_url, href)
                        doc = self._scrape_act(act_url)
                        if doc:
                            documents.append(doc)
                            count += 1
                            time.sleep(1)  # Be respectful
                    except Exception as e:
                        logger.warning(f"Failed to scrape act {href}: {e}")
                        continue
                        
        except Exception as e:
            logger.error(f"Error scraping IndiaCode: {e}")
            
        return documents
    
    def _scrape_act(self, act_url: str) -> Optional[ScrapedDocument]:
        """Scrape individual act"""
        try:
            response = self.session.get(act_url)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Extract title
            title_elem = soup.find('h1') or soup.find('title')
            title = title_elem.get_text().strip() if title_elem else "Unknown Act"
            
            # Extract content
            content_elem = soup.find('div', class_='content') or soup.find('main')
            if not content_elem:
                content_elem = soup.find('body')
            
            content = content_elem.get_text(separator='\n').strip() if content_elem else ""
            
            # Clean content
            content = self._clean_text(content)
            
            if len(content) < 100:  # Skip if too short
                return None
                
            return ScrapedDocument(
                title=title,
                content=content,
                source_url=act_url,
                source_type='indiacode',
                doc_type='statute',
                metadata={'scraped_at': datetime.now().isoformat()}
            )
            
        except Exception as e:
            logger.warning(f"Failed to scrape act {act_url}: {e}")
            return None
    
    def _clean_text(self, text: str) -> str:
        """Clean scraped text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove common headers/footers
        text = re.sub(r'(Page \d+ of \d+|©.*?Government of India)', '', text)
        return text.strip()

class ConstitutionScraper:
    """Scraper for Constitution of India"""
    
    def __init__(self):
        self.base_url = "https://legislative.gov.in/constitution-of-india"
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })
    
    def scrape_constitution(self) -> List[ScrapedDocument]:
        """Scrape Constitution of India"""
        documents = []
        
        try:
            response = self.session.get(self.base_url)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Extract constitution content
            content_elem = soup.find('div', class_='content') or soup.find('main')
            if not content_elem:
                content_elem = soup.find('body')
            
            content = content_elem.get_text(separator='\n').strip() if content_elem else ""
            content = self._clean_text(content)
            
            if content:
                documents.append(ScrapedDocument(
                    title="Constitution of India",
                    content=content,
                    source_url=self.base_url,
                    source_type='constitution',
                    doc_type='constitution',
                    metadata={'scraped_at': datetime.now().isoformat()}
                ))
                
        except Exception as e:
            logger.error(f"Error scraping Constitution: {e}")
            
        return documents
    
    def _clean_text(self, text: str) -> str:
        """Clean constitution text"""
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

class GitHubScraper:
    """Scraper for GitHub Indian Law JSON repositories"""
    
    def __init__(self):
        self.repos = [
            "lexglue/legal_case_reports",
            "LexGLUE/legal_case_reports", 
            "indian-legal-ner/indian-legal-ner",
            "legal-nlp/indian-legal-datasets"
        ]
    
    def scrape_repos(self) -> List[ScrapedDocument]:
        """Scrape GitHub repositories for legal data"""
        documents = []
        
        for repo in self.repos:
            try:
                # This is a simplified approach - in practice, you'd use GitHub API
                logger.info(f"Scraping GitHub repo: {repo}")
                # Implementation would use GitHub API to fetch JSON files
                # For now, return empty list
                pass
            except Exception as e:
                logger.warning(f"Failed to scrape repo {repo}: {e}")
                
        return documents

class ZenodoScraper:
    """Scraper for Zenodo datasets"""
    
    def __init__(self):
        self.zenodo_id = "4065965"  # Annotated Central Acts Dataset
    
    def scrape_zenodo(self) -> List[ScrapedDocument]:
        """Scrape Zenodo dataset"""
        documents = []
        
        try:
            # Zenodo API endpoint
            api_url = f"https://zenodo.org/api/records/{self.zenodo_id}"
            response = requests.get(api_url)
            response.raise_for_status()
            
            data = response.json()
            
            # Get download links
            files = data.get('files', [])
            for file_info in files:
                if file_info.get('type') == 'json':
                    # Download and process JSON file
                    file_url = file_info['links']['self']
                    file_docs = self._process_zenodo_file(file_url)
                    documents.extend(file_docs)
                    
        except Exception as e:
            logger.error(f"Error scraping Zenodo: {e}")
            
        return documents
    
    def _process_zenodo_file(self, file_url: str) -> List[ScrapedDocument]:
        """Process individual Zenodo file"""
        documents = []
        
        try:
            response = requests.get(file_url)
            response.raise_for_status()
            
            data = response.json()
            
            # Process based on file structure
            if isinstance(data, list):
                for item in data:
                    doc = self._create_document_from_zenodo_item(item)
                    if doc:
                        documents.append(doc)
            elif isinstance(data, dict):
                doc = self._create_document_from_zenodo_item(data)
                if doc:
                    documents.append(doc)
                    
        except Exception as e:
            logger.warning(f"Failed to process Zenodo file {file_url}: {e}")
            
        return documents
    
    def _create_document_from_zenodo_item(self, item: Dict[str, Any]) -> Optional[ScrapedDocument]:
        """Create document from Zenodo item"""
        try:
            title = item.get('title', 'Unknown Document')
            content = item.get('content', '') or item.get('text', '')
            
            if not content or len(content) < 50:
                return None
                
            return ScrapedDocument(
                title=title,
                content=content,
                source_url=item.get('url', ''),
                source_type='zenodo',
                doc_type=item.get('type', 'document'),
                section_article=item.get('section'),
                date=item.get('date'),
                metadata=item.get('metadata', {})
            )
            
        except Exception as e:
            logger.warning(f"Failed to create document from Zenodo item: {e}")
            return None

class ILDCScraper:
    """Scraper for ILDC Corpus"""
    
    def __init__(self):
        self.ildc_url = "https://paperswithcode.com/dataset/ildc"
    
    def scrape_ildc(self) -> List[ScrapedDocument]:
        """Scrape ILDC corpus"""
        documents = []
        
        try:
            # ILDC is typically available as a dataset
            # This would need to be downloaded from the official source
            logger.info("ILDC scraping not implemented - requires dataset download")
            # Implementation would download and process ILDC dataset
            
        except Exception as e:
            logger.error(f"Error scraping ILDC: {e}")
            
        return documents

class KaggleScraper:
    """Scraper for Kaggle legal datasets"""
    
    def __init__(self):
        self.datasets = [
            "abhinavralhan/indian-legal-documents",
            "abhinavralhan/indian-constitution",
            "abhinavralhan/supreme-court-cases"
        ]
    
    def scrape_kaggle(self) -> List[ScrapedDocument]:
        """Scrape Kaggle datasets"""
        documents = []
        
        try:
            # Set up Kaggle API
            kaggle.api.authenticate()
            
            for dataset in self.datasets:
                try:
                    # Download dataset
                    kaggle.api.dataset_download_files(dataset, path=f"data/kaggle/{dataset}", unzip=True)
                    
                    # Process downloaded files
                    dataset_path = Path(f"data/kaggle/{dataset}")
                    if dataset_path.exists():
                        file_docs = self._process_kaggle_dataset(dataset_path)
                        documents.extend(file_docs)
                        
                except Exception as e:
                    logger.warning(f"Failed to download Kaggle dataset {dataset}: {e}")
                    continue
                    
        except Exception as e:
            logger.error(f"Error setting up Kaggle API: {e}")
            
        return documents
    
    def _process_kaggle_dataset(self, dataset_path: Path) -> List[ScrapedDocument]:
        """Process downloaded Kaggle dataset"""
        documents = []
        
        for file_path in dataset_path.rglob("*.json"):
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                if isinstance(data, list):
                    for item in data:
                        doc = self._create_document_from_kaggle_item(item, str(file_path))
                        if doc:
                            documents.append(doc)
                elif isinstance(data, dict):
                    doc = self._create_document_from_kaggle_item(data, str(file_path))
                    if doc:
                        documents.append(doc)
                        
            except Exception as e:
                logger.warning(f"Failed to process Kaggle file {file_path}: {e}")
                
        return documents
    
    def _create_document_from_kaggle_item(self, item: Dict[str, Any], file_path: str) -> Optional[ScrapedDocument]:
        """Create document from Kaggle item"""
        try:
            title = item.get('title', item.get('case_name', 'Unknown Document'))
            content = item.get('content', item.get('text', item.get('judgment', '')))
            
            if not content or len(content) < 50:
                return None
                
            return ScrapedDocument(
                title=title,
                content=content,
                source_url=item.get('url', ''),
                source_type='kaggle',
                doc_type=item.get('type', 'document'),
                section_article=item.get('section'),
                date=item.get('date'),
                metadata={'file_path': file_path, **item.get('metadata', {})}
            )
            
        except Exception as e:
            logger.warning(f"Failed to create document from Kaggle item: {e}")
            return None

class HuggingFaceScraper:
    """Scraper for Hugging Face legal datasets"""
    
    def __init__(self):
        self.datasets = [
            ("lex_glue", "scotus"),  # Supreme Court cases
            ("lex_glue", "case_hold"),  # Case holdings
            ("lex_glue", "ecthr_a"),  # European Court of Human Rights
        ]
    
    def scrape_huggingface(self) -> List[ScrapedDocument]:
        """Scrape Hugging Face datasets"""
        documents = []
        
        for dataset_name, config_name in self.datasets:
            try:
                logger.info(f"Loading Hugging Face dataset: {dataset_name}/{config_name}")
                dataset = load_dataset(dataset_name, config_name)
                
                # Process each split
                for split_name, split_data in dataset.items():
                    for i, item in enumerate(split_data):
                        doc = self._create_document_from_hf_item(item, f"{dataset_name}/{config_name}", split_name)
                        if doc:
                            documents.append(doc)
                        
                        # Limit to prevent too many documents
                        if i >= 50:  # Limit to 50 documents per split
                            break
                            
            except Exception as e:
                logger.warning(f"Failed to load Hugging Face dataset {dataset_name}/{config_name}: {e}")
                continue
                
        return documents
    
    def _create_document_from_hf_item(self, item: Dict[str, Any], dataset_name: str, split_name: str) -> Optional[ScrapedDocument]:
        """Create document from Hugging Face item"""
        try:
            title = item.get('title', item.get('case_name', f"{dataset_name}_{split_name}"))
            content = item.get('text', item.get('content', item.get('judgment', '')))
            
            if not content or len(content) < 50:
                return None
                
            return ScrapedDocument(
                title=title,
                content=content,
                source_url=item.get('url', ''),
                source_type='huggingface',
                doc_type=item.get('type', 'document'),
                section_article=item.get('section'),
                date=item.get('date'),
                metadata={'dataset': dataset_name, 'split': split_name, **item.get('metadata', {})}
            )
            
        except Exception as e:
            logger.warning(f"Failed to create document from Hugging Face item: {e}")
            return None

class LawSumScraper:
    """Scraper for LawSum and Legal NLP corpora"""
    
    def __init__(self):
        self.sources = [
            "https://huggingface.co/datasets/legal_summarization",
            "https://arxiv.org/abs/2106.15599"  # LawSum paper
        ]
    
    def scrape_lawsum(self) -> List[ScrapedDocument]:
        """Scrape LawSum and related corpora"""
        documents = []
        
        try:
            # Try Hugging Face first
            hf_scraper = HuggingFaceScraper()
            hf_docs = hf_scraper.scrape_huggingface()
            documents.extend(hf_docs)
            
            # Additional LawSum specific processing could go here
            
        except Exception as e:
            logger.error(f"Error scraping LawSum: {e}")
            
        return documents

class LegalDataScraper:
    """Main scraper orchestrator"""
    
    def __init__(self):
        self.scrapers = {
            'indiacode': IndiaCodeScraper(),
            'constitution': ConstitutionScraper(),
            'github': GitHubScraper(),
            'zenodo': ZenodoScraper(),
            'ildc': ILDCScraper(),
            'kaggle': KaggleScraper(),
            'huggingface': HuggingFaceScraper(),
            'lawsum': LawSumScraper()
        }
    
    def scrape_all_sources(self, sources: Optional[List[str]] = None) -> List[ScrapedDocument]:
        """Scrape all or specified sources"""
        all_documents = []
        
        sources_to_scrape = sources or list(self.scrapers.keys())
        
        for source in sources_to_scrape:
            if source not in self.scrapers:
                logger.warning(f"Unknown source: {source}")
                continue
                
            try:
                logger.info(f"Scraping {source}...")
                scraper = self.scrapers[source]
                
                if source == 'indiacode':
                    docs = scraper.scrape_acts(max_acts=20)  # Limit for demo
                elif source == 'constitution':
                    docs = scraper.scrape_constitution()
                elif source == 'github':
                    docs = scraper.scrape_repos()
                elif source == 'zenodo':
                    docs = scraper.scrape_zenodo()
                elif source == 'ildc':
                    docs = scraper.scrape_ildc()
                elif source == 'kaggle':
                    docs = scraper.scrape_kaggle()
                elif source == 'huggingface':
                    docs = scraper.scrape_huggingface()
                elif source == 'lawsum':
                    docs = scraper.scrape_lawsum()
                else:
                    continue
                
                logger.info(f"Scraped {len(docs)} documents from {source}")
                all_documents.extend(docs)
                
            except Exception as e:
                logger.error(f"Error scraping {source}: {e}")
                continue
                
        return all_documents
    
    def save_scraped_data(self, documents: List[ScrapedDocument], output_file: str = "data/scraped_legal_data.json"):
        """Save scraped data to JSON file"""
        try:
            os.makedirs(os.path.dirname(output_file), exist_ok=True)
            
            data = []
            for doc in documents:
                data.append({
                    'title': doc.title,
                    'content': doc.content,
                    'source_url': doc.source_url,
                    'source_type': doc.source_type,
                    'section_article': doc.section_article,
                    'date': doc.date,
                    'doc_type': doc.doc_type,
                    'metadata': doc.metadata
                })
            
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
                
            logger.info(f"Saved {len(documents)} documents to {output_file}")
            
        except Exception as e:
            logger.error(f"Error saving scraped data: {e}")

def main():
    """Main function for testing scraper"""
    scraper = LegalDataScraper()
    
    # Scrape specific sources
    sources = ['constitution', 'huggingface']  # Start with these for testing
    documents = scraper.scrape_all_sources(sources)
    
    print(f"Scraped {len(documents)} documents")
    
    # Save to file
    scraper.save_scraped_data(documents)
    
    # Print sample
    for i, doc in enumerate(documents[:3]):
        print(f"\n--- Document {i+1} ---")
        print(f"Title: {doc.title}")
        print(f"Source: {doc.source_type}")
        print(f"Content: {doc.content[:200]}...")

if __name__ == "__main__":
    main()
