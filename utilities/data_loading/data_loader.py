"""
Data Loader Module for Agentic Legal RAG
Handles ingestion of various legal data sources including ILDC, InLegalBERT, Bare Acts, and Indian Kanoon
"""

import os
import json
import zipfile
import requests
import pandas as pd
from typing import List, Dict, Any, Optional, Generator
from pathlib import Path
import logging
from dataclasses import dataclass
import time
import re
from bs4 import BeautifulSoup
import PyPDF2
import docx
from io import BytesIO
import subprocess
import sys

from config import config

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class LegalDocument:
    """Standardized legal document structure"""
    doc_id: str
    title: str
    content: str
    doc_type: str  # 'statute', 'judgment', 'amendment', 'constitution'
    source: str
    url: Optional[str] = None
    court: Optional[str] = None
    date: Optional[str] = None
    citations: List[str] = None
    metadata: Dict[str, Any] = None
    
    def __post_init__(self):
        if self.citations is None:
            self.citations = []
        if self.metadata is None:
            self.metadata = {}

class ILDCLoader:
    """Loader for Indian Legal Documents Corpus (ILDC)"""
    
    def __init__(self, data_dir: str = None):
        self.data_dir = data_dir or config.database.ILDC_DIR
        self.ensure_data_dir()
    
    def ensure_data_dir(self):
        """Create data directory if it doesn't exist"""
        os.makedirs(self.data_dir, exist_ok=True)
    
    def load_from_zip(self, zip_path: str) -> List[LegalDocument]:
        """Load ILDC data from zip file"""
        documents = []
        
        try:
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                # Extract to temporary directory
                extract_path = os.path.join(self.data_dir, "extracted_ildc")
                zip_ref.extractall(extract_path)
                
                # Process extracted files
                for root, dirs, files in os.walk(extract_path):
                    for file in files:
                        if file.endswith('.json'):
                            file_path = os.path.join(root, file)
                            docs = self._process_ildc_json(file_path)
                            documents.extend(docs)
                            
        except Exception as e:
            logger.error(f"Error loading ILDC from zip: {e}")
            
        return documents
    
    def _process_ildc_json(self, json_path: str) -> List[LegalDocument]:
        """Process individual ILDC JSON file"""
        documents = []
        
        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                
            if isinstance(data, list):
                for item in data:
                    doc = self._create_document_from_ildc(item)
                    if doc:
                        documents.append(doc)
            elif isinstance(data, dict):
                doc = self._create_document_from_ildc(data)
                if doc:
                    documents.append(doc)
                    
        except Exception as e:
            logger.error(f"Error processing ILDC JSON {json_path}: {e}")
            
        return documents
    
    def _create_document_from_ildc(self, item: Dict[str, Any]) -> Optional[LegalDocument]:
        """Create LegalDocument from ILDC item"""
        try:
            # Extract basic information
            doc_id = item.get('id', f"ildc_{hash(str(item))}")
            title = item.get('title', 'Untitled Document')
            content = item.get('text', '')
            
            # Determine document type
            doc_type = 'judgment'  # ILDC primarily contains judgments
            if 'statute' in title.lower() or 'act' in title.lower():
                doc_type = 'statute'
            
            # Extract metadata
            metadata = {
                'source_file': item.get('source_file', ''),
                'year': item.get('year', ''),
                'court': item.get('court', ''),
                'bench': item.get('bench', ''),
                'petitioner': item.get('petitioner', ''),
                'respondent': item.get('respondent', ''),
                'citations': item.get('citations', [])
            }
            
            return LegalDocument(
                doc_id=doc_id,
                title=title,
                content=content,
                doc_type=doc_type,
                source='ILDC',
                court=item.get('court', ''),
                date=item.get('date', ''),
                citations=item.get('citations', []),
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Error creating document from ILDC item: {e}")
            return None

class InLegalBERTLoader:
    """Loader for InLegalBERT legal statutes"""
    
    def __init__(self, data_dir: str = None):
        self.data_dir = data_dir or config.database.DATA_DIR
        self.ensure_data_dir()
    
    def ensure_data_dir(self):
        """Create data directory if it doesn't exist"""
        os.makedirs(self.data_dir, exist_ok=True)
    
    def load_from_huggingface(self) -> List[LegalDocument]:
        """Load InLegalBERT data from Hugging Face"""
        documents = []
        
        try:
            from datasets import load_dataset
            
            # Load InLegalBERT dataset
            dataset = load_dataset("law-ai/InLegalBERT", split="train")
            
            for item in dataset:
                doc = self._create_document_from_inlegalbert(item)
                if doc:
                    documents.append(doc)
                    
        except Exception as e:
            logger.error(f"Error loading InLegalBERT from Hugging Face: {e}")
            logger.info("Falling back to manual download instructions")
            self._provide_manual_download_instructions()
            
        return documents
    
    def _create_document_from_inlegalbert(self, item: Dict[str, Any]) -> Optional[LegalDocument]:
        """Create LegalDocument from InLegalBERT item"""
        try:
            doc_id = f"inlegalbert_{item.get('id', hash(str(item)))}"
            title = item.get('title', 'Legal Statute')
            content = item.get('text', '')
            
            return LegalDocument(
                doc_id=doc_id,
                title=title,
                content=content,
                doc_type='statute',
                source='InLegalBERT',
                metadata={
                    'section': item.get('section', ''),
                    'act': item.get('act', ''),
                    'year': item.get('year', ''),
                    'category': item.get('category', '')
                }
            )
            
        except Exception as e:
            logger.error(f"Error creating document from InLegalBERT item: {e}")
            return None
    
    def _provide_manual_download_instructions(self):
        """Provide instructions for manual download"""
        instructions = """
        Manual Download Instructions for InLegalBERT:
        1. Visit: https://huggingface.co/datasets/law-ai/InLegalBERT
        2. Download the dataset files
        3. Place them in: {}/inlegalbert/
        4. Run the loader again
        """.format(self.data_dir)
        logger.info(instructions)

class BareActsLoader:
    """Loader for Bare Acts (Government legal documents)"""
    
    def __init__(self, data_dir: str = None):
        self.data_dir = data_dir or config.database.BARE_ACTS_DIR
        self.ensure_data_dir()
    
    def ensure_data_dir(self):
        """Create data directory if it doesn't exist"""
        os.makedirs(self.data_dir, exist_ok=True)
    
    def load_from_directory(self, directory_path: str = None) -> List[LegalDocument]:
        """Load Bare Acts from directory containing PDF/DOCX files"""
        documents = []
        directory_path = directory_path or self.data_dir
        
        if not os.path.exists(directory_path):
            logger.warning(f"Bare Acts directory not found: {directory_path}")
            return documents
        
        for file_path in Path(directory_path).rglob("*"):
            if file_path.suffix.lower() in ['.pdf', '.docx', '.txt']:
                docs = self._process_file(file_path)
                documents.extend(docs)
                
        return documents
    
    def _process_file(self, file_path: Path) -> List[LegalDocument]:
        """Process individual file"""
        documents = []
        
        try:
            if file_path.suffix.lower() == '.pdf':
                docs = self._process_pdf(file_path)
            elif file_path.suffix.lower() == '.docx':
                docs = self._process_docx(file_path)
            elif file_path.suffix.lower() == '.txt':
                docs = self._process_txt(file_path)
            else:
                return documents
                
            documents.extend(docs)
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            
        return documents
    
    def _process_pdf(self, file_path: Path) -> List[LegalDocument]:
        """Process PDF file"""
        documents = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                content = ""
                
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
                
                # Extract metadata
                title = file_path.stem
                doc_id = f"bare_act_{hash(str(file_path))}"
                
                # Determine act type from filename
                doc_type = 'statute'
                if 'amendment' in title.lower():
                    doc_type = 'amendment'
                elif 'constitution' in title.lower():
                    doc_type = 'constitution'
                
                doc = LegalDocument(
                    doc_id=doc_id,
                    title=title,
                    content=content,
                    doc_type=doc_type,
                    source='Bare Acts',
                    metadata={
                        'file_path': str(file_path),
                        'file_size': file_path.stat().st_size,
                        'pages': len(pdf_reader.pages)
                    }
                )
                documents.append(doc)
                
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            
        return documents
    
    def _process_docx(self, file_path: Path) -> List[LegalDocument]:
        """Process DOCX file"""
        documents = []
        
        try:
            doc = docx.Document(file_path)
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            title = file_path.stem
            doc_id = f"bare_act_{hash(str(file_path))}"
            
            doc = LegalDocument(
                doc_id=doc_id,
                title=title,
                content=content,
                doc_type='statute',
                source='Bare Acts',
                metadata={
                    'file_path': str(file_path),
                    'file_size': file_path.stat().st_size
                }
            )
            documents.append(doc)
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            
        return documents
    
    def _process_txt(self, file_path: Path) -> List[LegalDocument]:
        """Process TXT file"""
        documents = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            title = file_path.stem
            doc_id = f"bare_act_{hash(str(file_path))}"
            
            doc = LegalDocument(
                doc_id=doc_id,
                title=title,
                content=content,
                doc_type='statute',
                source='Bare Acts',
                metadata={
                    'file_path': str(file_path),
                    'file_size': file_path.stat().st_size
                }
            )
            documents.append(doc)
            
        except Exception as e:
            logger.error(f"Error processing TXT {file_path}: {e}")
            
        return documents

class KaggleLoader:
    """Loader for Kaggle legal datasets"""
    
    def __init__(self, data_dir: str = None):
        self.data_dir = data_dir or config.database.DATA_DIR
        self.kaggle_dir = os.path.join(self.data_dir, "kaggle")
        self.ensure_data_dir()
        self._check_kaggle_api()
    
    def ensure_data_dir(self):
        """Create data directory if it doesn't exist"""
        os.makedirs(self.kaggle_dir, exist_ok=True)
    
    def _check_kaggle_api(self):
        """Check if Kaggle API is installed and configured"""
        try:
            import kaggle
            logger.info("Kaggle API is available")
            return True
        except ImportError:
            logger.warning("Kaggle API not installed. Installing...")
            try:
                subprocess.check_call([sys.executable, "-m", "pip", "install", "kaggle"])
                logger.info("Kaggle API installed successfully")
                return True
            except subprocess.CalledProcessError:
                logger.error("Failed to install Kaggle API")
                return False
    
    def download_dataset(self, dataset_name: str, unzip: bool = True) -> str:
        """Download dataset from Kaggle"""
        try:
            from kaggle.api.kaggle_api_extended import KaggleApi
            
            api = KaggleApi()
            api.authenticate()
            
            # Download dataset
            download_path = os.path.join(self.kaggle_dir, dataset_name.replace("/", "_"))
            api.dataset_download_files(
                dataset_name, 
                path=download_path, 
                unzip=unzip
            )
            
            logger.info(f"Downloaded dataset {dataset_name} to {download_path}")
            return download_path
            
        except Exception as e:
            logger.error(f"Error downloading dataset {dataset_name}: {e}")
            return None
    
    def load_legal_dataset(self, dataset_name: str) -> List[LegalDocument]:
        """Load legal dataset from Kaggle"""
        documents = []
        
        # Download dataset
        dataset_path = self.download_dataset(dataset_name)
        if not dataset_path:
            return documents
        
        # Process downloaded files
        for root, dirs, files in os.walk(dataset_path):
            for file in files:
                file_path = os.path.join(root, file)
                docs = self._process_kaggle_file(file_path, dataset_name)
                documents.extend(docs)
        
        logger.info(f"Loaded {len(documents)} documents from Kaggle dataset: {dataset_name}")
        return documents
    
    def _process_kaggle_file(self, file_path: str, dataset_name: str) -> List[LegalDocument]:
        """Process individual file from Kaggle dataset"""
        documents = []
        
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.json':
                docs = self._process_json_file(file_path, dataset_name)
            elif file_ext == '.csv':
                docs = self._process_csv_file(file_path, dataset_name)
            elif file_ext == '.txt':
                docs = self._process_txt_file(file_path, dataset_name)
            elif file_ext == '.pdf':
                docs = self._process_pdf_file(file_path, dataset_name)
            else:
                logger.warning(f"Unsupported file type: {file_ext}")
                return documents
            
            documents.extend(docs)
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
        
        return documents
    
    def _process_json_file(self, file_path: str, dataset_name: str) -> List[LegalDocument]:
        """Process JSON file from Kaggle"""
        documents = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            if isinstance(data, list):
                for item in data:
                    doc = self._create_document_from_kaggle_json(item, dataset_name)
                    if doc:
                        documents.append(doc)
            elif isinstance(data, dict):
                doc = self._create_document_from_kaggle_json(data, dataset_name)
                if doc:
                    documents.append(doc)
                    
        except Exception as e:
            logger.error(f"Error processing JSON file {file_path}: {e}")
        
        return documents
    
    def _process_csv_file(self, file_path: str, dataset_name: str) -> List[LegalDocument]:
        """Process CSV file from Kaggle"""
        documents = []
        
        try:
            df = pd.read_csv(file_path)
            
            for index, row in df.iterrows():
                doc = self._create_document_from_kaggle_csv(row, dataset_name, index)
                if doc:
                    documents.append(doc)
                    
        except Exception as e:
            logger.error(f"Error processing CSV file {file_path}: {e}")
        
        return documents
    
    def _process_txt_file(self, file_path: str, dataset_name: str) -> List[LegalDocument]:
        """Process TXT file from Kaggle"""
        documents = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Split content into chunks if it's too long
            chunks = self._split_text_into_chunks(content, 2000)
            
            for i, chunk in enumerate(chunks):
                doc = LegalDocument(
                    doc_id=f"kaggle_{dataset_name}_{os.path.basename(file_path)}_{i}",
                    title=f"{os.path.basename(file_path)} - Part {i+1}",
                    content=chunk,
                    doc_type='legal_document',
                    source=f'Kaggle - {dataset_name}',
                    metadata={
                        'file_path': file_path,
                        'chunk_number': i,
                        'total_chunks': len(chunks)
                    }
                )
                documents.append(doc)
                
        except Exception as e:
            logger.error(f"Error processing TXT file {file_path}: {e}")
        
        return documents
    
    def _process_pdf_file(self, file_path: str, dataset_name: str) -> List[LegalDocument]:
        """Process PDF file from Kaggle"""
        documents = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                content = ""
                
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
                
                # Split content into chunks
                chunks = self._split_text_into_chunks(content, 2000)
                
                for i, chunk in enumerate(chunks):
                    doc = LegalDocument(
                        doc_id=f"kaggle_{dataset_name}_{os.path.basename(file_path)}_{i}",
                        title=f"{os.path.basename(file_path)} - Part {i+1}",
                        content=chunk,
                        doc_type='legal_document',
                        source=f'Kaggle - {dataset_name}',
                        metadata={
                            'file_path': file_path,
                            'chunk_number': i,
                            'total_chunks': len(chunks),
                            'pages': len(pdf_reader.pages)
                        }
                    )
                    documents.append(doc)
                    
        except Exception as e:
            logger.error(f"Error processing PDF file {file_path}: {e}")
        
        return documents
    
    def _create_document_from_kaggle_json(self, item: Dict[str, Any], dataset_name: str) -> Optional[LegalDocument]:
        """Create LegalDocument from Kaggle JSON item"""
        try:
            # Extract common fields
            doc_id = item.get('id', f"kaggle_{hash(str(item))}")
            title = item.get('title', item.get('case_title', item.get('name', 'Untitled Document')))
            content = item.get('text', item.get('content', item.get('description', '')))
            
            # Determine document type based on content
            doc_type = 'legal_document'
            if 'judgment' in title.lower() or 'case' in title.lower():
                doc_type = 'judgment'
            elif 'act' in title.lower() or 'statute' in title.lower():
                doc_type = 'statute'
            elif 'constitution' in title.lower():
                doc_type = 'constitution'
            
            return LegalDocument(
                doc_id=f"kaggle_{dataset_name}_{doc_id}",
                title=title,
                content=content,
                doc_type=doc_type,
                source=f'Kaggle - {dataset_name}',
                court=item.get('court', ''),
                date=item.get('date', item.get('year', '')),
                citations=item.get('citations', []),
                metadata={
                    'dataset': dataset_name,
                    'original_id': doc_id,
                    'raw_data': item
                }
            )
            
        except Exception as e:
            logger.error(f"Error creating document from Kaggle JSON: {e}")
            return None
    
    def _create_document_from_kaggle_csv(self, row: pd.Series, dataset_name: str, index: int) -> Optional[LegalDocument]:
        """Create LegalDocument from Kaggle CSV row"""
        try:
            # Convert row to dict
            row_dict = row.to_dict()
            
            # Extract text content from common column names
            content = ""
            for col in ['text', 'content', 'description', 'case_text', 'judgment']:
                if col in row_dict and pd.notna(row_dict[col]):
                    content += str(row_dict[col]) + " "
            
            if not content.strip():
                return None
            
            # Extract title
            title = ""
            for col in ['title', 'case_title', 'name', 'case_name']:
                if col in row_dict and pd.notna(row_dict[col]):
                    title = str(row_dict[col])
                    break
            
            if not title:
                title = f"Document {index + 1}"
            
            return LegalDocument(
                doc_id=f"kaggle_{dataset_name}_{index}",
                title=title,
                content=content.strip(),
                doc_type='legal_document',
                source=f'Kaggle - {dataset_name}',
                metadata={
                    'dataset': dataset_name,
                    'row_index': index,
                    'raw_data': row_dict
                }
            )
            
        except Exception as e:
            logger.error(f"Error creating document from Kaggle CSV: {e}")
            return None
    
    def _split_text_into_chunks(self, text: str, chunk_size: int = 2000) -> List[str]:
        """Split text into chunks for processing"""
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        words = text.split()
        current_chunk = []
        current_length = 0
        
        for word in words:
            if current_length + len(word) + 1 > chunk_size and current_chunk:
                chunks.append(' '.join(current_chunk))
                current_chunk = [word]
                current_length = len(word)
            else:
                current_chunk.append(word)
                current_length += len(word) + 1
        
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        return chunks

class IndianKanoonLoader:
    """Loader for Indian Kanoon API data"""
    
    def __init__(self, api_key: str = None):
        self.api_key = api_key or config.api.INDIAN_KANOON_API_KEY
        self.base_url = config.api.KANOON_BASE_URL
        self.rate_limit_delay = 60 / config.api.API_RATE_LIMIT  # seconds between requests
    
    def load_judgments(self, query: str, max_results: int = 100) -> List[LegalDocument]:
        """Load judgments from Indian Kanoon API"""
        documents = []
        
        if not self.api_key or self.api_key == "YOUR_KANOON_API_KEY_HERE":
            logger.warning("Indian Kanoon API key not provided. Skipping API data loading.")
            return documents
        
        try:
            # Implement API calls to Indian Kanoon
            # This is a placeholder - actual implementation would depend on their API
            logger.info(f"Loading judgments for query: {query}")
            
            # Simulate API call with rate limiting
            time.sleep(self.rate_limit_delay)
            
            # Placeholder response processing
            # In real implementation, this would parse actual API responses
            
        except Exception as e:
            logger.error(f"Error loading from Indian Kanoon API: {e}")
            
        return documents
    
    def load_statutes(self, act_name: str) -> List[LegalDocument]:
        """Load specific act from Indian Kanoon"""
        documents = []
        
        if not self.api_key or self.api_key == "YOUR_KANOON_API_KEY_HERE":
            logger.warning("Indian Kanoon API key not provided. Skipping API data loading.")
            return documents
        
        try:
            logger.info(f"Loading statute: {act_name}")
            # Implement statute loading from API
            time.sleep(self.rate_limit_delay)
            
        except Exception as e:
            logger.error(f"Error loading statute {act_name}: {e}")
            
        return documents

class DataLoader:
    """Main data loader orchestrator"""
    
    def __init__(self):
        self.ildc_loader = ILDCLoader()
        self.inlegalbert_loader = InLegalBERTLoader()
        self.bare_acts_loader = BareActsLoader()
        self.kanoon_loader = IndianKanoonLoader()
        self.kaggle_loader = KaggleLoader()
    
    def load_all_sources(self) -> List[LegalDocument]:
        """Load data from all available sources"""
        all_documents = []
        
        logger.info("Starting data loading from all sources...")
        
        # Load from ILDC
        logger.info("Loading ILDC data...")
        ildc_docs = self._load_ildc_data()
        all_documents.extend(ildc_docs)
        logger.info(f"Loaded {len(ildc_docs)} documents from ILDC")
        
        # Load from InLegalBERT
        logger.info("Loading InLegalBERT data...")
        inlegalbert_docs = self.ildc_loader.load_from_huggingface()
        all_documents.extend(inlegalbert_docs)
        logger.info(f"Loaded {len(inlegalbert_docs)} documents from InLegalBERT")
        
        # Load from Bare Acts
        logger.info("Loading Bare Acts data...")
        bare_acts_docs = self.bare_acts_loader.load_from_directory()
        all_documents.extend(bare_acts_docs)
        logger.info(f"Loaded {len(bare_acts_docs)} documents from Bare Acts")
        
        # Load from Indian Kanoon (if API key available)
        logger.info("Loading Indian Kanoon data...")
        kanoon_docs = self._load_kanoon_data()
        all_documents.extend(kanoon_docs)
        logger.info(f"Loaded {len(kanoon_docs)} documents from Indian Kanoon")
        
        logger.info(f"Total documents loaded: {len(all_documents)}")
        return all_documents
    
    def load_kaggle_dataset(self, dataset_name: str) -> List[LegalDocument]:
        """Load specific Kaggle dataset"""
        logger.info(f"Loading Kaggle dataset: {dataset_name}")
        documents = self.kaggle_loader.load_legal_dataset(dataset_name)
        logger.info(f"Loaded {len(documents)} documents from Kaggle dataset: {dataset_name}")
        return documents
    
    def load_kaggle_datasets(self, dataset_names: List[str]) -> List[LegalDocument]:
        """Load multiple Kaggle datasets"""
        all_documents = []
        
        for dataset_name in dataset_names:
            docs = self.load_kaggle_dataset(dataset_name)
            all_documents.extend(docs)
        
        logger.info(f"Total documents loaded from Kaggle: {len(all_documents)}")
        return all_documents
    
    def _load_ildc_data(self) -> List[LegalDocument]:
        """Load ILDC data with fallback options"""
        # Try to find ILDC zip file in data directory
        ildc_zip = None
        for file in os.listdir(self.ildc_loader.data_dir):
            if file.endswith('.zip') and 'ildc' in file.lower():
                ildc_zip = os.path.join(self.ildc_loader.data_dir, file)
                break
        
        if ildc_zip and os.path.exists(ildc_zip):
            return self.ildc_loader.load_from_zip(ildc_zip)
        else:
            logger.warning("ILDC zip file not found. Please download and place in data/ildc/")
            return []
    
    def _load_kanoon_data(self) -> List[LegalDocument]:
        """Load Indian Kanoon data with sample queries"""
        documents = []
        
        # Sample queries for testing
        sample_queries = [
            "Article 21 Constitution of India",
            "Section 302 IPC",
            "Hindu Marriage Act",
            "Consumer Protection Act"
        ]
        
        for query in sample_queries:
            docs = self.kanoon_loader.load_judgments(query, max_results=10)
            documents.extend(docs)
        
        return documents
    
    def save_documents(self, documents: List[LegalDocument], output_path: str):
        """Save documents to JSON file for later processing"""
        try:
            data = []
            for doc in documents:
                data.append({
                    'doc_id': doc.doc_id,
                    'title': doc.title,
                    'content': doc.content,
                    'doc_type': doc.doc_type,
                    'source': doc.source,
                    'url': doc.url,
                    'court': doc.court,
                    'date': doc.date,
                    'citations': doc.citations,
                    'metadata': doc.metadata
                })
            
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            
            logger.info(f"Saved {len(documents)} documents to {output_path}")
            
        except Exception as e:
            logger.error(f"Error saving documents: {e}")

def main():
    """Main function for testing data loading"""
    loader = DataLoader()
    
    # Load all data
    documents = loader.load_all_sources()
    
    # Save to file
    output_path = os.path.join(config.database.DATA_DIR, "loaded_documents.json")
    loader.save_documents(documents, output_path)
    
    # Print summary
    print(f"Loaded {len(documents)} documents")
    
    # Group by source
    sources = {}
    for doc in documents:
        source = doc.source
        if source not in sources:
            sources[source] = 0
        sources[source] += 1
    
    print("Documents by source:")
    for source, count in sources.items():
        print(f"  {source}: {count}")

if __name__ == "__main__":
    main()

